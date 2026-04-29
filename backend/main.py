import os
import io
import uuid
from pathlib import Path
from datetime import datetime, timezone

import httpx
import boto3
import numpy as np
from dotenv import load_dotenv
from pypdf import PdfReader
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from botocore.exceptions import ClientError
from botocore.config import Config
import pypdfium2 as pdfium

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from rapidocr_onnxruntime import RapidOCR
except Exception:
    RapidOCR = None

# Load environment variables from .env
load_dotenv()

PORT = int(os.getenv("PORT", "5000"))
FLASK_SERVICE_URL = os.getenv("FLASK_SERVICE_URL", "http://localhost:5001")
MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT", "localhost:9000")
MINIO_ACCESS_KEY = os.getenv("MINIO_ACCESS_KEY", "")
MINIO_SECRET_KEY = os.getenv("MINIO_SECRET_KEY", "")
MINIO_BUCKET_NAME = os.getenv("MINIO_BUCKET_NAME", "resumes")
MINIO_USE_SSL = os.getenv("MINIO_USE_SSL", "False").lower() == "true"

OCR_ENGINE = RapidOCR() if RapidOCR is not None else None

app = FastAPI(title="Cloud Resume Analyzer API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# Initialize MinIO client for resume file uploads.
def init_minio_client():
    try:
        if not MINIO_ACCESS_KEY or not MINIO_SECRET_KEY or not MINIO_ENDPOINT:
            print("[WARN] MinIO env vars missing. File upload will be disabled.")
            return None

        client = boto3.client(
            "s3",
            endpoint_url=f"{'https' if MINIO_USE_SSL else 'http'}://{MINIO_ENDPOINT}",
            aws_access_key_id=MINIO_ACCESS_KEY,
            aws_secret_access_key=MINIO_SECRET_KEY,
            use_ssl=MINIO_USE_SSL,
            config=Config(connect_timeout=5, read_timeout=10, retries={"max_attempts": 2}),
        )
        print(f"[INFO] MinIO client initialized at {MINIO_ENDPOINT}")
        return client
    except Exception as exc:
        print(f"[ERROR] MinIO init failed: {exc}")
        return None


minio_client = init_minio_client()


# Ensure the target MinIO bucket exists before uploading.
async def ensure_minio_bucket():
    try:
        if minio_client is None:
            return

        print(f"[INFO] Checking MinIO bucket '{MINIO_BUCKET_NAME}'...")
        minio_client.head_bucket(Bucket=MINIO_BUCKET_NAME)
        print(f"[INFO] Bucket '{MINIO_BUCKET_NAME}' already exists.")
    except ClientError as exc:
        error_code = str(exc.response.get("Error", {}).get("Code", ""))
        if error_code in {"404", "NoSuchBucket", "NotFound"}:
            print(f"[INFO] Creating bucket '{MINIO_BUCKET_NAME}'...")
            minio_client.create_bucket(Bucket=MINIO_BUCKET_NAME)
            print(f"[INFO] Bucket '{MINIO_BUCKET_NAME}' created successfully.")
        else:
            print(f"[WARN] Bucket check failed: {exc}")
    except Exception as exc:
        print(f"[WARN] MinIO bucket check skipped: {exc}")


def ocr_pdf_text(file_bytes: bytes, page_limit: int = 3) -> str:
    if OCR_ENGINE is None:
        return ""

    try:
        pdf = pdfium.PdfDocument(io.BytesIO(file_bytes))
        page_count = min(len(pdf), page_limit)
        page_texts = []

        for index in range(page_count):
            page = pdf[index]
            bitmap = page.render(scale=2)
            image = np.array(bitmap.to_pil())
            ocr_result, _ = OCR_ENGINE(image)
            lines = [item[1].strip() for item in (ocr_result or []) if len(item) > 1 and item[1].strip()]
            if lines:
                page_texts.append("\n".join(lines))

        text = "\n".join(page_texts)
        print(f"[INFO] OCR extracted {len(text)} characters from PDF")
        return text
    except Exception as exc:
        print(f"[WARN] OCR fallback failed: {exc}")
        return ""


# Parse uploaded resume file into plain text for downstream ML analysis.
async def parse_resume(file: UploadFile, file_bytes: bytes) -> str:
    try:
        name = (file.filename or "resume.txt").lower()

        if name.endswith(".txt"):
            print("[INFO] Parsing TXT resume.")
            text = file_bytes.decode("utf-8", errors="ignore")
            print(f"[INFO] Extracted {len(text)} characters from TXT")
            return text

        if name.endswith(".pdf"):
            print("[INFO] Parsing PDF resume.")
            reader = PdfReader(io.BytesIO(file_bytes))
            print(f"[INFO] PDF has {len(reader.pages)} pages")
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages)
            if len(text.strip()) < 50 and pdfplumber is not None:
                print("[WARN] Pypdf extracted little text. Trying pdfplumber fallback.")
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    fallback_pages = [(page.extract_text() or "") for page in pdf.pages]
                fallback_text = "\n".join(fallback_pages)
                if len(fallback_text.strip()) > len(text.strip()):
                    text = fallback_text
            if len(text.strip()) < 50:
                print("[WARN] PDF appears to be image-only. Trying OCR fallback.")
                ocr_text = ocr_pdf_text(file_bytes)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
            print(f"[INFO] Extracted {len(text)} characters from PDF")
            return text

        if name.endswith(".docx"):
            print("[INFO] Parsing DOCX resume.")
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            text = "\n".join(paragraphs)
            print(f"[INFO] Extracted {len(text)} characters from DOCX ({len(paragraphs)} paragraphs)")
            return text

        raise HTTPException(status_code=400, detail="Unsupported file type. Use TXT, PDF, or DOCX.")
    except HTTPException:
        raise
    except Exception as exc:
        print(f"[ERROR] Resume parse failed: {exc}")
        raise HTTPException(status_code=500, detail="Failed to parse resume file.")


# Upload the original resume file to MinIO and return object metadata.
async def upload_to_minio(file: UploadFile, file_bytes: bytes):
    try:
        if minio_client is None:
            print("[WARN] MinIO unavailable. Returning local placeholder metadata.")
            key = f"{uuid.uuid4()}-{file.filename}"
            return {
                "minioKey": key,
                "minioUrl": f"local://{key}",
            }

        await ensure_minio_bucket()
        key = f"{uuid.uuid4()}-{file.filename}"
        print(f"[INFO] Uploading file to MinIO: {key}")
        minio_client.put_object(
            Bucket=MINIO_BUCKET_NAME,
            Key=key,
            Body=file_bytes,
            ContentType=file.content_type or "application/octet-stream",
        )

        minio_url = f"{'https' if MINIO_USE_SSL else 'http'}://{MINIO_ENDPOINT}/{MINIO_BUCKET_NAME}/{key}"
        print(f"[INFO] Upload successful: {minio_url}")
        return {
            "minioKey": key,
            "minioUrl": minio_url,
        }
    except Exception as exc:
        print(f"[WARN] MinIO upload failed, falling back to local metadata: {exc}")
        key = f"{uuid.uuid4()}-{file.filename}"
        # Persist the file locally so we don't lose the uploaded bytes.
        try:
            uploads_dir = Path(__file__).parent / "uploads"
            uploads_dir.mkdir(parents=True, exist_ok=True)
            local_path = uploads_dir / key
            with local_path.open("wb") as fh:
                fh.write(file_bytes)
            file_url = f"file://{local_path.resolve()}"
            print(f"[INFO] Saved uploaded file locally to {local_path}")
            return {
                "minioKey": key,
                "minioUrl": file_url,
            }
        except Exception as write_exc:
            print(f"[ERROR] Failed to persist file locally: {write_exc}")
            return {
                "minioKey": key,
                "minioUrl": f"local://{key}",
            }


# Call Flask ML microservice to compute semantic score and entity extraction.
async def call_ml_service(resume_text: str, job_description: str):
    last_error = None
    try:
        print("[INFO] Calling ML service for analysis.")
        async with httpx.AsyncClient(timeout=httpx.Timeout(45.0, connect=5.0)) as client:
            for attempt in range(1, 4):
                try:
                    response = await client.post(
                        f"{FLASK_SERVICE_URL}/analyze",
                        json={
                            "resumeText": resume_text,
                            "jobDescription": job_description,
                        },
                    )
                    response.raise_for_status()
                    print("[INFO] ML service response received.")
                    return response.json()
                except httpx.HTTPStatusError as exc:
                    last_error = exc
                    status_code = exc.response.status_code
                    body = exc.response.text[:500]
                    print(f"[WARN] ML service returned {status_code} on attempt {attempt}: {body}")
                    if status_code < 500 or attempt == 3:
                        raise HTTPException(
                            status_code=502,
                            detail=f"ML service returned {status_code}.",
                        )
                except (httpx.ConnectError, httpx.ReadTimeout, httpx.RemoteProtocolError) as exc:
                    last_error = exc
                    print(f"[WARN] ML service attempt {attempt} failed: {exc}")
                    if attempt == 3:
                        raise HTTPException(status_code=502, detail="Failed to connect to ML service.")
    except HTTPException:
        raise
    except httpx.HTTPError as exc:
        last_error = exc
        print(f"[ERROR] ML service call failed: {exc}")
        raise HTTPException(status_code=502, detail="Failed to connect to ML service.")
    finally:
        if last_error is not None:
            print(f"[INFO] ML service last error: {last_error}")





@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "service": "cloud-resume-analyzer-fastapi",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "storage": {
            "provider": "minio",
            "endpoint": MINIO_ENDPOINT,
            "bucket": MINIO_BUCKET_NAME,
        },
    }


@app.post("/api/analyze")
async def analyze_resume(
    resume: UploadFile = File(...),
    jobDescription: str = Form(default=""),
    title: str = Form(default=""),
    company: str = Form(default=""),
):
    try:
        print("[INFO] Analyze request received.")
        file_bytes = await resume.read()
        if not file_bytes:
            raise HTTPException(status_code=400, detail="Uploaded file is empty.")

        parsed_text = await parse_resume(resume, file_bytes)
        if not parsed_text.strip():
            raise HTTPException(status_code=400, detail="Unable to extract text from resume.")

        uploaded = await upload_to_minio(resume, file_bytes)
        ml_result = await call_ml_service(parsed_text, jobDescription)

        resume_id = str(uuid.uuid4())
        uploaded_at = datetime.now(timezone.utc).isoformat()
        analysis = {
            "score": ml_result.get("score", 0),
            "overallScore": ml_result.get("score", 0),
            "matchedKeywords": ml_result.get("matchedKeywords", []),
            "missingKeywords": ml_result.get("missingKeywords", []),
            "skillGaps": ml_result.get("skillGaps", []),
            "recommendations": ml_result.get("recommendations", []),
            "roleRecommendations": ml_result.get("roleRecommendations", []),
            "resumeIssues": ml_result.get("resumeIssues", []),
            "strengthSignals": ml_result.get("strengthSignals", []),
            "improvementPlan": ml_result.get("improvementPlan", []),
        }

        print("[INFO] Analyze flow completed.")
        return {
            "message": "Resume analyzed successfully.",
            "resumeId": resume_id,
            "uploadedAt": uploaded_at,
            "file": {
                "originalName": resume.filename,
                "minioKey": uploaded["minioKey"],
                "minioUrl": uploaded["minioUrl"],
                "fileType": resume.content_type or "application/octet-stream",
                "fileSizeKb": round(len(file_bytes) / 1024, 2),
            },
            "analysis": analysis,
            "extracted": ml_result.get(
                "extracted",
                {
                    "name": "",
                    "email": "",
                    "phone": "",
                    "skills": [],
                    "education": [],
                    "experience": [],
                    "totalExperienceYears": 0,
                },
            ),
        }
    except HTTPException:
        raise
    except Exception as exc:
        print(f"[ERROR] Unexpected analyze failure: {exc}")
        raise HTTPException(status_code=500, detail="Unexpected server error.")


@app.get("/api/history")
async def history():
    print("[INFO] History endpoint called. Local file storage mode - no persistent history available.")
    return []


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("main:app", host="0.0.0.0", port=PORT, reload=True)
